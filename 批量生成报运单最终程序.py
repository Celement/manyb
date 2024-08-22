from openpyxl import load_workbook
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from mailmerge import MailMerge
import openpyxl
import os
import pandas as pd
import qrcode

# pyinstaller -F -w -i=lg.ico czwin.py  --hidden-import=openpyxl.cell._writer --onefile

def gen_doc_text(tempath,savename,param_key):

    print("开始生成word分件")
    doc = MailMerge(tempath)

    # 将内容添加到Word模板文件中 参数名与Word模板中的域名相同
    doc.merge(
      ordernumber= str(param_key["ordernumber"]),
      reference= str(param_key["reference"]),
      weight= str(param_key["weight"]),
      store= str(param_key["store"]),
      pickupaddress= str(param_key["pickupaddress"]),
      ofcases= str(param_key["ofcases"]),
      totalpallets= str(param_key["totalpallets"]),
    )

    doc.write(savename)
def gen_code(pathname,text):
    print("二维码")
    # 传入将要生成二维码的URL
    img = qrcode.make(text)
    # 保存
    img.save("./code/"+pathname+".png")
def addPicToWord(picPath,docment):
    '''
    图片插入到word文档中
    :param picPath: 图片的路径
    :param docment: word文档对象
    :return:
    '''
    for table in docment.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 根据文档中的占位符定位图片插入的位置
                    if '<<pic>>' in paragraph.text:
                        print("找到了占位符")
                        # 把占位符去掉
                        paragraph.text = paragraph.text.replace('<<pic>>', '')
                        run = paragraph.add_run('')
                        if picPath!=None:
                            run.add_break()
                            # 添加图片并根据表格的高度进行缩放
                            run.add_picture(picPath, width=Inches(0.7),height=Inches(0.7))
                            # source_height=picture.height
                            # afterheight=row.height
                            # print(afterheight)
                            # picture.height =int(afterheight)
                            # picture.width = int(picture.width * (afterheight / source_height))

def unmerge_cell(excel_name, sheet_name,cf_name):
    # 打开工作簿并获取sheet
    wb = openpyxl.load_workbook(excel_name)
    sheet = wb[sheet_name]

    merged_cells = list(sheet.merged_cells.ranges)

    # 遍历每个合并单元格
    for merged_cell in merged_cells.copy():
        # 获取合并单元格的值
        merged_value = sheet.cell(row=merged_cell.min_row, column=merged_cell.min_col).value

        # 对'合并单元格'进行拆分
        sheet.unmerge_cells(str(merged_cell))

        # 将值分配给每个单元格
        for row in range(merged_cell.min_row, merged_cell.max_row + 1):
            for column in range(merged_cell.min_col, merged_cell.max_col + 1):
                cell = sheet.cell(row=row, column=column)
                cell.value = merged_value

    # 保存工作簿
    wb.save(cf_name)

def gen_byd(doc_template_name,excel_data_name,sheet_name,excel_data_template_name):

    unmerge_cell(excel_data_name, sheet_name, excel_data_template_name)
    # 读取Excel文件
    df = pd.read_excel(excel_data_template_name)

    # print(df.head())
    for index, row in df.iterrows():
        param_key = {
            "ordernumber": row[0],
            "reference": row[1],
            "weight": row[2],
            "store": row[3],
            "pickupaddress": row[4],
            "ofcases": row[5],
            "totalpallets": row[6],
        }
        print(row[0])
        if isinstance(row[0], str):
            print("开始生成二维码")
            if not os.path.exists("./code/"):
                os.makedirs("./code/")
            if not os.path.exists("./报运/"):
                os.makedirs("./报运/")
            if not os.path.exists("./最终/"):
                os.makedirs("./最终/")
            # 生成二维码
            pathname = str(row[0])
            gen_code(pathname, str(row[7]))
            print("生成二维码成功")

            gen_doc_text(doc_template_name, "./报运/" + row[0] + ".docx", param_key)
            print("单个word生成成功")
            print("开始插入图片")
            doc = Document("./报运/" + row[0] + ".docx")

            addPicToWord("./code/" + row[0] + ".png", doc)

            doc.save("./最终/" + row[0] + ".docx")
            print("图片插入成功")





if __name__ == '__main__':
    # #原始excel
    # excel_data_name= 'May 10 list.xlsx'
    # # 拆分后excel
    # excel_data_template_name = "temp.xlsx"
    # # 工作区
    # sheet_name = 'Sheet1'
    # # 模版文件
    # doc_template_name = "BOL模板.docx"

    excel_data_name=input("请输入excel文件路径:")
    sheet_name = input("请输入你要操作的工作区:")

    excel_data_template_name=input("请给拆分单元格后的excel取个名字:")
    doc_template_name=input("请输入模版文件的名字:")

    gen_byd(excel_data_name=excel_data_name,sheet_name=sheet_name,doc_template_name=doc_template_name,excel_data_template_name=excel_data_template_name)









